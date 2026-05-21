from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def exportar_relatorio_word(df, periodo_analisado="Período não informado"):
    outputs_dir = Path("outputs")
    outputs_dir.mkdir(exist_ok=True)

    caminho_saida = outputs_dir / "relatorio_tecnico_ava.docx"

    doc = Document()

    estilo = doc.styles["Normal"]
    estilo.font.name = "Arial"
    estilo.font.size = Pt(10)

    # CAPA
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ATHENA SCIENTIFIC\n")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 51, 102)

    run = p.add_run("AVA INTELLIGENCE\n")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("RELATÓRIO TÉCNICO DE MONITORAMENTO ACADÊMICO\n").bold = True
    p.add_run(f"Período analisado: {periodo_analisado}\n")
    p.add_run(f"Data da análise: {datetime.now().strftime('%d/%m/%Y')}")

    doc.add_page_break()

    # INTRODUÇÃO
    doc.add_heading("1. INTRODUÇÃO", level=1)
    doc.add_paragraph(
        "O presente relatório técnico apresenta os resultados do monitoramento acadêmico "
        "realizado pelo sistema AVA Intelligence, desenvolvido no ecossistema Athena Scientific. "
        "A ferramenta utiliza princípios de Learning Analytics para acompanhar o engajamento "
        "discente no Ambiente Virtual de Aprendizagem, subsidiando ações de gestão acadêmica, "
        "coordenação de curso e Núcleo Docente Estruturante."
    )

    # METODOLOGIA
    doc.add_heading("2. METODOLOGIA", level=1)
    doc.add_paragraph(
        "A análise foi realizada a partir do cruzamento entre os dados extraídos dos arquivos "
        "institucionais em PDF, contendo matrícula e nome dos estudantes, e os registros de acesso "
        "extraídos do Moodle. O ranking de engajamento foi construído com base no número de acessos, "
        "último acesso registrado e dias sem acesso."
    )

    doc.add_paragraph(
        "Foram adotados os seguintes critérios de risco acadêmico: até 3 dias sem acesso, sem risco; "
        "de 4 a 7 dias sem acesso, risco médio; mais de 7 dias sem acesso, risco elevado."
    )

    # RESULTADOS
    doc.add_heading("3. RESULTADOS", level=1)

    total_alunos = len(df)

    media_acessos = df["acessos"].mean() if "acessos" in df.columns else 0

    sem_acesso = len(df[df["acessos"] == 0]) if "acessos" in df.columns else 0

    risco_medio = len(df[df["risco"].astype(str).str.lower().str.contains("médio|medio", na=False)]) if "risco" in df.columns else 0
    risco_elevado = len(df[df["risco"].astype(str).str.lower().str.contains("elevado", na=False)]) if "risco" in df.columns else 0

    indicadores = [
        ("Total de alunos analisados", total_alunos),
        ("Média de acessos", round(media_acessos, 2)),
        ("Alunos sem acesso", sem_acesso),
        ("Alunos em risco médio", risco_medio),
        ("Alunos em risco elevado", risco_elevado),
    ]

    tabela_ind = doc.add_table(rows=1, cols=2)
    tabela_ind.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela_ind.style = "Table Grid"

    hdr = tabela_ind.rows[0].cells
    hdr[0].text = "Indicador"
    hdr[1].text = "Resultado"

    for indicador, valor in indicadores:
        row = tabela_ind.add_row().cells
        row[0].text = str(indicador)
        row[1].text = str(valor)

    # TABELA COMPLETA
    doc.add_heading("4. TABELA COMPLETA DE MONITORAMENTO", level=1)

    colunas = [
        "matricula",
        "nome",
        "periodo",
        "turma",
        "acessos",
        "ultimo_acesso",
        "dias_sem_acesso",
        "ranking",
        "risco",
        "parecer_ia",
    ]

    colunas_existentes = [c for c in colunas if c in df.columns]

    tabela = doc.add_table(rows=1, cols=len(colunas_existentes))
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.style = "Table Grid"

    header = tabela.rows[0].cells
    for i, col in enumerate(colunas_existentes):
        header[i].text = col.replace("_", " ").upper()

    for _, linha in df[colunas_existentes].iterrows():
        cells = tabela.add_row().cells
        for i, col in enumerate(colunas_existentes):
            valor = linha[col]
            cells[i].text = "" if str(valor) == "nan" else str(valor)

    # ANÁLISE PEDAGÓGICA
    doc.add_heading("5. ANÁLISE PEDAGÓGICA AUTOMÁTICA", level=1)

    doc.add_paragraph(
        f"A análise identificou {total_alunos} estudantes monitorados no período avaliado. "
        f"A média geral de acessos foi de {round(media_acessos, 2)}, com {sem_acesso} estudantes "
        f"sem registro de acesso. Foram identificados {risco_medio} estudantes em risco médio e "
        f"{risco_elevado} estudantes em risco elevado. Esses dados indicam a necessidade de "
        "acompanhamento contínuo, contato ativo com os estudantes em maior vulnerabilidade acadêmica "
        "e utilização dos indicadores de engajamento como apoio à gestão pedagógica."
    )

    # CONCLUSÃO
    doc.add_heading("6. CONCLUSÃO", level=1)

    doc.add_paragraph(
        "O relatório produzido pelo AVA Intelligence permite à coordenação de curso, ao NDE e à gestão "
        "acadêmica acompanhar de forma sistemática o engajamento discente no Ambiente Virtual de "
        "Aprendizagem. A automatização do processo fortalece a tomada de decisão baseada em evidências "
        "e contribui para ações preventivas de acompanhamento pedagógico."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ATHENA SCIENTIFIC — AVA INTELLIGENCE").bold = True

    doc.save(caminho_saida)

    return str(caminho_saida)
