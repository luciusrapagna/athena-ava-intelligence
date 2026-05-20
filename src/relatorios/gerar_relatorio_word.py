from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from datetime import datetime
import pandas as pd


def configurar_fonte_padrao(doc):
    estilo = doc.styles["Normal"]
    fonte = estilo.font
    fonte.name = "Times New Roman"
    fonte.size = Pt(12)


def formatar_paragrafo(paragrafo, negrito=False, italico=False, tamanho=12):
    for run in paragrafo.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(tamanho)
        run.bold = negrito
        run.italic = italico


def aplicar_borda_celula(celula):
    tc = celula._tc
    tcPr = tc.get_or_add_tcPr()

    bordas = tcPr.first_child_found_in("w:tcBorders")

    if bordas is None:
        bordas = OxmlElement("w:tcBorders")
        tcPr.append(bordas)

    for lado in ["top", "bottom"]:
        borda = bordas.find(qn(f"w:{lado}"))

        if borda is None:
            borda = OxmlElement(f"w:{lado}")
            bordas.append(borda)

        borda.set(qn("w:val"), "single")
        borda.set(qn("w:sz"), "6")
        borda.set(qn("w:space"), "0")
        borda.set(qn("w:color"), "000000")

    for lado in ["left", "right"]:
        borda = bordas.find(qn(f"w:{lado}"))

        if borda is None:
            borda = OxmlElement(f"w:{lado}")
            bordas.append(borda)

        borda.set(qn("w:val"), "nil")


def aplicar_sombreamento(celula, cor="F2F2F2"):
    tcPr = celula._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), cor)
    tcPr.append(shd)


def formatar_celula(celula, negrito=False, tamanho=8):
    celula.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for paragrafo in celula.paragraphs:
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for run in paragrafo.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(tamanho)
            run.bold = negrito
            run.font.color.rgb = RGBColor(0, 0, 0)


def criar_tabela_padrao(doc, titulo, dataframe, colunas, fonte, limite_linhas=None):
    titulo_paragrafo = doc.add_paragraph()
    titulo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = titulo_paragrafo.add_run(titulo)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    if limite_linhas is not None:
        dataframe = dataframe.head(limite_linhas)

    tabela = doc.add_table(
        rows=1,
        cols=len(colunas)
    )

    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = True

    cabecalho = tabela.rows[0].cells

    for i, coluna in enumerate(colunas):
        cabecalho[i].text = coluna
        formatar_celula(cabecalho[i], negrito=True, tamanho=8)
        aplicar_borda_celula(cabecalho[i])

    for indice, (_, linha) in enumerate(dataframe.iterrows()):
        row = tabela.add_row().cells

        for i, coluna in enumerate(colunas):
            row[i].text = str(linha.get(coluna, ""))
            formatar_celula(row[i], tamanho=7)
            aplicar_borda_celula(row[i])

            if indice % 2 == 0:
                aplicar_sombreamento(row[i], "F2F2F2")

    fonte_paragrafo = doc.add_paragraph()
    fonte_paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run_fonte = fonte_paragrafo.add_run(fonte)
    run_fonte.italic = True
    run_fonte.font.name = "Times New Roman"
    run_fonte.font.size = Pt(10)

    doc.add_paragraph("")


def gerar_relatorio_word(df, dias_analise):
    data_hoje = datetime.now().strftime("%d-%m-%Y_%H-%M-%S")

    nome_relatorio = (
        f"outputs/word/relatorio_monitoramento_{data_hoje}.docx"
    )

    doc = Document()
    configurar_fonte_padrao(doc)

    titulo = doc.add_heading(
        f"ATHENA AVA INTELLIGENCE - {data_hoje}",
        level=0
    )
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading(
        "Relatório Institucional de Monitoramento Acadêmico",
        level=1
    )

    p_data = doc.add_paragraph(
        f"Período analisado: últimos {dias_analise} dias."
    )
    formatar_paragrafo(p_data)

    doc.add_heading("1. INTRODUÇÃO METODOLÓGICA", level=1)

    p_intro = doc.add_paragraph(
        f"""
O presente relatório foi elaborado pelo sistema Athena AVA Intelligence com o objetivo de monitorar indicadores de engajamento acadêmico no Ambiente Virtual de Aprendizagem institucional.

A análise considerou registros de acesso estudantil referentes aos últimos {dias_analise} dias, permitindo identificar padrões de participação, frequência digital, ranking de engajamento e potenciais situações de risco acadêmico.

A classificação adotada considerou os seguintes critérios: estudantes com até 4 dias sem acesso foram classificados como sem risco; estudantes com 5 a 6 dias sem acesso foram classificados como médio risco; e estudantes com 7 dias ou mais sem atividade foram classificados como alto risco.

O ranking de engajamento foi construído com base no total de acessos, na existência de acesso registrado no período e no número de dias desde o último acesso ao AVA.
"""
    )
    formatar_paragrafo(p_intro)

    df["Risco"] = df["Risco"].astype(str)
    df["Nível de engajamento"] = df["Nível de engajamento"].astype(str)

    df_risco = df[
        df["Risco"].str.contains("ALTO|MÉDIO", na=False, regex=True)
    ].copy()

    df_sem_risco = df[
        ~df["Risco"].str.contains("ALTO|MÉDIO", na=False, regex=True)
    ].copy()

    df_ranking = df.sort_values(
        by=["Ranking"],
        ascending=True
    ).copy()

    total = len(df)
    total_risco = len(df_risco)
    total_sem_risco = len(df_sem_risco)
    alto = len(df[df["Risco"].str.contains("ALTO", na=False)])
    medio = len(df[df["Risco"].str.contains("MÉDIO", na=False)])

    alto_eng = len(
        df[df["Nível de engajamento"].str.contains("Alto", na=False)]
    )

    moderado_eng = len(
        df[df["Nível de engajamento"].str.contains("moderado", na=False)]
    )

    baixo_eng = len(
        df[df["Nível de engajamento"].str.contains("Baixo", na=False)]
    )

    df_resumo = pd.DataFrame(
        {
            "Indicador": [
                "Total de alunos analisados",
                "Alunos em risco",
                "Alunos sem risco",
                "Alto risco",
                "Médio risco",
                "Alto engajamento",
                "Engajamento moderado",
                "Baixo engajamento",
            ],
            "Número": [
                total,
                total_risco,
                total_sem_risco,
                alto,
                medio,
                alto_eng,
                moderado_eng,
                baixo_eng,
            ],
            "Percentual (%)": [
                100,
                round((total_risco / total) * 100, 2) if total > 0 else 0,
                round((total_sem_risco / total) * 100, 2) if total > 0 else 0,
                round((alto / total) * 100, 2) if total > 0 else 0,
                round((medio / total) * 100, 2) if total > 0 else 0,
                round((alto_eng / total) * 100, 2) if total > 0 else 0,
                round((moderado_eng / total) * 100, 2) if total > 0 else 0,
                round((baixo_eng / total) * 100, 2) if total > 0 else 0,
            ],
        }
    )

    doc.add_heading("2. RESUMO EXECUTIVO", level=1)

    criar_tabela_padrao(
        doc=doc,
        titulo="Tabela 1 – Distribuição dos estudantes segundo risco e engajamento no AVA",
        dataframe=df_resumo,
        colunas=["Indicador", "Número", "Percentual (%)"],
        fonte="Fonte: Elaboração própria a partir dos registros de acesso do Ambiente Virtual de Aprendizagem."
    )

    doc.add_heading("3. RANKING DE ENGAJAMENTO", level=1)

    p_rank = doc.add_paragraph(
        "A tabela a seguir apresenta o ranking de engajamento dos estudantes, considerando acesso registrado, frequência de participação e recência do último acesso ao AVA."
    )
    formatar_paragrafo(p_rank)

    colunas_ranking = [
        "Ranking",
        "Matrícula",
        "Aluno",
        "Período",
        "Turma",
        "Total de acessos",
        "Dias sem acesso",
        "Nível de engajamento",
    ]

    criar_tabela_padrao(
        doc=doc,
        titulo="Tabela 2 – Ranking de engajamento dos estudantes no AVA",
        dataframe=df_ranking,
        colunas=colunas_ranking,
        fonte="Fonte: Elaboração própria a partir da lista oficial de alunos e dos logs do AVA."
    )

    doc.add_heading("4. ESTUDANTES EM RISCO", level=1)

    p_risco = doc.add_paragraph(
        "A tabela a seguir apresenta os estudantes classificados em médio ou alto risco, considerando o tempo sem acesso ao AVA e a necessidade de acompanhamento pedagógico."
    )
    formatar_paragrafo(p_risco)

    colunas_risco = [
        "Matrícula",
        "Aluno",
        "Período",
        "Turma",
        "Entrou no AVA",
        "Total de acessos",
        "Dias sem acesso",
        "Risco",
        "Sugestão de ação",
    ]

    criar_tabela_padrao(
        doc=doc,
        titulo="Tabela 3 – Estudantes em risco e sugestões de ação pedagógica",
        dataframe=df_risco,
        colunas=colunas_risco,
        fonte="Fonte: Elaboração própria a partir dos critérios institucionais de risco acadêmico digital."
    )

    doc.add_heading("5. ESTUDANTES SEM RISCO", level=1)

    p_sem_risco = doc.add_paragraph(
        "A tabela a seguir apresenta os estudantes que não foram classificados em risco no período analisado, indicando manutenção do acompanhamento regular."
    )
    formatar_paragrafo(p_sem_risco)

    colunas_sem_risco = [
        "Matrícula",
        "Aluno",
        "Período",
        "Turma",
        "Entrou no AVA",
        "Total de acessos",
        "Dias sem acesso",
        "Nível de engajamento",
    ]

    criar_tabela_padrao(
        doc=doc,
        titulo="Tabela 4 – Estudantes sem risco acadêmico no período analisado",
        dataframe=df_sem_risco,
        colunas=colunas_sem_risco,
        fonte="Fonte: Elaboração própria a partir da lista oficial de alunos e dos logs do AVA."
    )

    doc.add_heading("6. PARECER PEDAGÓGICO AUTOMÁTICO", level=1)

    p_parecer_intro = doc.add_paragraph(
        "A tabela a seguir apresenta pareceres pedagógicos automáticos gerados a partir dos indicadores de engajamento, acesso e risco acadêmico digital."
    )
    formatar_paragrafo(p_parecer_intro)

    colunas_parecer = [
        "Aluno",
        "Período",
        "Turma",
        "Risco",
        "Nível de engajamento",
        "Parecer pedagógico",
    ]

    criar_tabela_padrao(
        doc=doc,
        titulo="Tabela 5 – Parecer pedagógico automático por estudante",
        dataframe=df,
        colunas=colunas_parecer,
        fonte="Fonte: Elaboração própria por meio do sistema Athena AVA Intelligence."
    )

    doc.add_heading("7. RECOMENDAÇÕES INSTITUCIONAIS", level=1)

    p_rec = doc.add_paragraph(
        """
Recomenda-se que os estudantes em alto risco sejam priorizados para contato ativo pela coordenação, tutoria ou equipe pedagógica. Para estudantes em médio risco, sugere-se acompanhamento preventivo e orientação para retomada da participação no AVA. Para estudantes com baixo engajamento, recomenda-se verificar dificuldades técnicas, problemas de login, vínculo no ambiente virtual ou ausência de engajamento acadêmico.
"""
    )
    formatar_paragrafo(p_rec)

    doc.add_heading("8. CONCLUSÃO INSTITUCIONAL", level=1)

    p_conclusao = doc.add_paragraph(
        """
O monitoramento automatizado do Ambiente Virtual de Aprendizagem permite identificar estudantes em diferentes níveis de risco acadêmico e engajamento digital, subsidiando ações institucionais de acompanhamento, permanência estudantil e intervenção pedagógica precoce.
"""
    )
    formatar_paragrafo(p_conclusao)

    doc.save(nome_relatorio)

    print("\nRelatório Word criado com sucesso.")
    print(nome_relatorio)